from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any
import json

def create_diagnosis_document(
    company_profile: Dict,
    edital_data: Dict,
    diagnosis_results: Dict,
    weighted_score: float,
    classification: str,
    recommendation: str
) -> Document:
    """
    Create a Word document with the complete diagnosis.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("DIAGNÓSTICO DE ADERÊNCIA - FINEP COPILOT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n").bold = True
    metadata.add_run(f"Empresa: {company_profile.get('nome_empresa', 'N/A')}\n")
    metadata.add_run(f"CNPJ: {company_profile.get('cnpj', 'N/A')}\n")
    metadata.add_run(f"Edital: {edital_data.get('titulo', 'N/A')}")
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    doc.add_heading("1. RESUMO EXECUTIVO", level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"Pontuação Geral: {weighted_score}/10\n").bold = True
    summary.add_run(f"Classificação: {classification}\n").bold = True
    summary.add_run(f"Recomendação: {recommendation}")
    
    doc.add_paragraph()
    
    # Binary Criteria
    doc.add_heading("2. CRITÉRIOS BINÁRIOS (ELIMINATÓRIOS)", level=1)
    
    binary_criteria = diagnosis_results.get("criterios_binarios", {})
    all_passed = True
    
    for criterion_name, criterion_data in binary_criteria.items():
        passed = criterion_data.get("passou", False)
        message = criterion_data.get("mensagem", "")
        
        status = "✓ PASSOU" if passed else "✗ NÃO PASSOU"
        color = RGBColor(0, 128, 0) if passed else RGBColor(255, 0, 0)
        
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{criterion_name.replace('_', ' ').title()}: {status}\n")
        run.font.color.rgb = color
        run.bold = True
        
        p.add_run(f"Detalhes: {message}")
        
        if not passed:
            all_passed = False
    
    if not all_passed:
        warning = doc.add_paragraph()
        warning_run = warning.add_run("⚠️ ATENÇÃO: Existem critérios binários não atendidos. A proposta será inabilitada.")
        warning_run.font.color.rgb = RGBColor(255, 0, 0)
        warning_run.bold = True
    
    doc.add_paragraph()
    
    # Graduated Criteria
    doc.add_heading("3. CRITÉRIOS GRADUAIS (PONTUAÇÃO)", level=1)
    
    graduated_criteria = diagnosis_results.get("criterios_graduais", {})
    weights = {
        "aderencia_tematica": 0.25,
        "maturidade_trl": 0.20,
        "capacidade_tecnica": 0.15,
        "historico_pd": 0.15,
        "contrapartida": 0.10,
        "impacto_escalabilidade": 0.10,
        "alinhamento_geografico_preferencial": 0.05
    }
    
    # Create table for scores
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "Critério"
    header_cells[1].text = "Peso"
    header_cells[2].text = "Nota"
    header_cells[3].text = "Justificativa"
    
    for criterion_name, criterion_data in graduated_criteria.items():
        score = criterion_data.get("nota", 0)
        justification = criterion_data.get("justificativa", "")
        weight = weights.get(criterion_name, 0)
        
        row_cells = table.add_row().cells
        row_cells[0].text = criterion_name.replace('_', ' ').title()
        row_cells[1].text = f"{weight*100:.0f}%"
        row_cells[2].text = f"{score}/10"
        row_cells[3].text = justification
    
    doc.add_paragraph()
    
    # Company Profile
    doc.add_heading("4. PERFIL DA EMPRESA", level=1)
    
    profile_items = [
        ("Nome", "nome_empresa"),
        ("CNPJ", "cnpj"),
        ("Natureza Jurídica", "natureza_juridica"),
        ("Porte", "porte"),
        ("Estado/Região", "estado_regiao"),
        ("Área de Atuação", "area_atuacao"),
        ("Descrição Técnica", "descricao_tecnica"),
        ("Histórico P&D", "historico_pd"),
        ("Certificações", "certificacoes"),
        ("Contrapartida Estimada", "contrapartida_percentual"),
        ("TRL Habitual", "trl_habitual")
    ]
    
    for label, key in profile_items:
        value = company_profile.get(key, "Não informado")
        p = doc.add_paragraph(f"{label}: {value}", style="List Bullet")
    
    doc.add_paragraph()
    
    # Edital Information
    doc.add_heading("5. INFORMAÇÕES DO EDITAL", level=1)
    
    edital_items = [
        ("Título", "titulo"),
        ("Objetivo", "objetivo"),
        ("Orçamento Total", "orcamento_total"),
        ("Valor Mínimo", "valor_minimo"),
        ("Valor Máximo", "valor_maximo"),
        ("Contrapartida Mínima", "contrapartida_minima"),
        ("TRL Mínimo", "trl_minimo"),
        ("TRL Máximo", "trl_maximo"),
        ("Prazo de Encerramento", "prazo_encerramento")
    ]
    
    for label, key in edital_items:
        value = edital_data.get(key, "Não informado")
        p = doc.add_paragraph(f"{label}: {value}", style="List Bullet")
    
    doc.add_paragraph()
    
    # Recommendations
    doc.add_heading("6. RECOMENDAÇÕES", level=1)
    
    recommendations = diagnosis_results.get("recomendacao_geral", "Sem recomendações específicas")
    doc.add_paragraph(recommendations)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("_" * 80)
    footer_text = doc.add_paragraph(
        "Este documento foi gerado automaticamente pelo FINEP Copilot. "
        "Para informações oficiais, consulte o edital completo no site da FINEP."
    )
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(9)
    footer_text.runs[0].font.italic = True
    
    return doc

def create_proposal_document(
    company_profile: Dict,
    edital_data: Dict,
    proposal_sections: Dict,
    project_idea: str
) -> Document:
    """
    Create a Word document with the proposal draft.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("PROPOSTA DE PROJETO - FINEP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n").bold = True
    metadata.add_run(f"Empresa: {company_profile.get('nome_empresa', 'N/A')}\n")
    metadata.add_run(f"CNPJ: {company_profile.get('cnpj', 'N/A')}\n")
    metadata.add_run(f"Edital: {edital_data.get('titulo', 'N/A')}\n")
    metadata.add_run(f"Ideia do Projeto: {project_idea}")
    
    doc.add_paragraph()
    
    # Proposal Sections
    sections_mapping = {
        "resumo_executivo": ("1. RESUMO EXECUTIVO", "Resumo Executivo"),
        "problema_oportunidade": ("2. PROBLEMA E OPORTUNIDADE", "Problema e Oportunidade"),
        "objetivo_geral": ("3. OBJETIVO GERAL", "Objetivo Geral"),
        "objetivos_especificos": ("4. OBJETIVOS ESPECÍFICOS", "Objetivos Específicos"),
        "metodologia": ("5. METODOLOGIA", "Metodologia"),
        "plano_trabalho": ("6. PLANO DE TRABALHO", "Plano de Trabalho"),
        "equipe_necessaria": ("7. PERFIL DA EQUIPE NECESSÁRIA", "Equipe Necessária"),
        "resultados_esperados": ("8. RESULTADOS ESPERADOS", "Resultados Esperados"),
        "indicadores_sucesso": ("9. INDICADORES DE SUCESSO", "Indicadores de Sucesso"),
        "impacto_economico_social": ("10. IMPACTO ECONÔMICO E SOCIAL", "Impacto Econômico e Social")
    }
    
    for key, (heading, _) in sections_mapping.items():
        doc.add_heading(heading, level=1)
        content = proposal_sections.get(key, "Conteúdo não disponível")
        doc.add_paragraph(content)
        doc.add_paragraph()
    
    # Notes
    doc.add_heading("NOTAS IMPORTANTES", level=1)
    notes = [
        "Este documento é um rascunho gerado automaticamente e deve ser revisado e adaptado.",
        "Consulte o edital completo para informações detalhadas sobre requisitos e critérios de avaliação.",
        "Certifique-se de que todas as informações estão alinhadas com o regulamento da FINEP.",
        "Inclua documentação comprobatória conforme exigido pelo edital.",
        "Revise cuidadosamente antes de submeter a proposta na plataforma FINEP."
    ]
    
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("_" * 80)
    footer_text = doc.add_paragraph(
        "Documento gerado pelo FINEP Copilot. "
        "Para informações oficiais, consulte: https://www.finep.gov.br"
    )
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(9)
    footer_text.runs[0].font.italic = True
    
    return doc

def save_document(doc: Document, filename: str) -> str:
    """
    Save document and return the file path.
    """
    doc.save(filename)
    return filename
